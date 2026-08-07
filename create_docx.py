# -*- coding: utf-8 -*-
"""
Create the final Word document with all 15 charts embedded.
Clean, visual-heavy layout for PPT use.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

CHART_DIR = r'C:\文件\大三\大三下\字里行间\新建文件夹\app\chart_images'
OUTPUT_PATH = r'C:\文件\大三\大三下\字里行间\新建文件夹\app\PPT流程图与数据图表.docx'

doc = Document()

# ---- Global styles ----
style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

for section in doc.sections:
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

# ---- Helper functions ----
def add_title(text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

def add_chart(image_name, width_inches=6.0):
    """Add a chart image centered on page."""
    path = os.path.join(CHART_DIR, image_name)
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=Inches(width_inches))
        doc.add_paragraph()  # small gap
    else:
        doc.add_paragraph(f'[Chart not found: {image_name}]')

def add_separator():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('─' * 50)
    run.font.size = Pt(6)
    run.font.color.rgb = RGBColor(0xA3, 0x9E, 0x93)


# ================================================================
# Cover
# ================================================================
add_title('LINES 字里行间', level=0)
add_title('PPT 流程图与数据图表', level=1)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run(
    '本文件包含 15 张可视化图表，涵盖两大核心系统：\n'
    '  • 六书驱动的汉字智能拆解引擎\n'
    '  • 多维汉字系联网络\n\n'
    '所有图表均为高清 PNG，可直接复制到 PPT 中使用。'
)
run.font.size = Pt(11)
run.font.name = 'Microsoft YaHei'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
doc.add_page_break()


# ================================================================
# Part 1: 拆解引擎
# ================================================================
add_title('第一部分：六书驱动的汉字智能拆解引擎', level=1)

add_title('1.1 系统四层架构', level=2)
add_chart('01_四层架构.png', 5.8)

add_title('1.2 六书分类与拆解策略', level=2)
add_chart('02_六书拆解策略.png', 5.8)

add_title('1.3 拆解处理流水线', level=2)
add_chart('03_拆解流水线.png', 6.0)

add_title('1.4 声旁可靠性三色评级', level=2)
add_chart('04_声旁三色评级.png', 5.8)

add_title('1.5 幽灵部件与变形偏旁总览', level=2)
add_chart('13_幽灵部件与变形偏旁.png', 6.0)

doc.add_page_break()

# ================================================================
# Part 2: 系联网络
# ================================================================
add_title('第二部分：多维汉字系联网络', level=1)

add_title('2.1 系统四阶段架构', level=2)
add_chart('05_系联网络架构.png', 5.8)

add_title('2.2 九种关系类型与四级优先级', level=2)
add_chart('06_九种关系类型.png', 5.8)

add_title('2.3 互斥去重机制', level=2)
add_chart('07_互斥去重机制.png', 5.8)

add_title('2.4 简繁智能去重（三通道）', level=2)
add_chart('14_简繁去重三通道.png', 5.8)

add_title('2.5 形音义三维亲密度评分', level=2)
add_chart('08_三维评分.png', 6.0)

add_title('2.6 七大倒排索引（知识图谱）', level=2)
add_chart('09_七大索引知识图谱.png', 5.8)

doc.add_page_break()

# ================================================================
# Part 3: 系统总览
# ================================================================
add_title('第三部分：系统关系总览', level=1)

add_title('3.1 两大系统定位对比', level=2)
add_chart('10_两大系统对比.png', 6.0)

add_title('3.2 端到端数据流全景', level=2)
add_chart('11_端到端数据流.png', 6.2)

add_title('3.3 系统数据架构总览', level=2)
add_chart('15_系统架构总览.png', 6.2)

add_title('3.4 布鲁姆认知目标 × 产品功能映射', level=2)
add_chart('12_布鲁姆认知映射.png', 6.0)


# ================================================================
# Save
# ================================================================
doc.save(OUTPUT_PATH)
print(f'Done! Document saved to: {OUTPUT_PATH}')
print(f'File size: {os.path.getsize(OUTPUT_PATH) / 1024:.1f} KB')
