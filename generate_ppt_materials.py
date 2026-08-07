# -*- coding: utf-8 -*-
"""
Generate PPT materials Word document with flowcharts and tables.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

OUTPUT_PATH = r'C:\文件\大三\大三下\字里行间\新建文件夹\app\PPT截图与数据参考.docx'

doc = Document()

# ============================================================
# Global styles
# ============================================================
style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

for section in doc.sections:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)


def set_cell_shading(cell, color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_cell_text(cell, text, bold=False, size=Pt(9), color=None, alignment=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = alignment
    run = p.add_run(text)
    run.bold = bold
    run.font.size = size
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)


def add_title(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    return heading


def add_para(doc, text, bold=False, size=Pt(10.5), alignment=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = alignment
    run = p.add_run(text)
    run.bold = bold
    run.font.size = size
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    return p


def make_table(doc, headers, rows, col_widths=None, header_color='2D5F8A'):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        add_cell_text(cell, header, bold=True, size=Pt(9), color=RGBColor(0xFF, 0xFF, 0xFF))
        set_cell_shading(cell, header_color)

    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            add_cell_text(cell, str(val), size=Pt(8.5))
            if r % 2 == 1:
                set_cell_shading(cell, 'F2F0EB')

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)

    doc.add_paragraph()
    return table


# ================================================================
# Cover page
# ================================================================
add_title(doc, 'LINES -- PPT Flowcharts & Data Reference', level=0)
add_para(doc, 'This document contains structured flowcharts (text descriptions for PPT SmartArt conversion) and tables for the two core systems of the LINES project.', size=Pt(11))
add_para(doc, 'Generated: 2026-06-03', size=Pt(9))
doc.add_page_break()

# ================================================================
# Part 1: Six-Book-Driven Chinese Character Intelligent Decomposition Engine
# ================================================================
add_title(doc, 'Part 1: Six-Book-Driven Chinese Character Intelligent Decomposition Engine', level=1)

# --- 1.1 Overall Architecture Flow ---
add_title(doc, '1.1 Overall Architecture (for PPT Vertical Process SmartArt)', level=2)
add_para(doc, '[PPT Tip: Use "Vertical Process" or "Segmented Process" SmartArt, organized as follows]', size=Pt(9), bold=True)

flow1_data = [
    ['Layer 1: Data Source', 'hanzi-dict.json (Six-Book classification + IDS notation + Shuowen Jiezi data)'],
    ['    |', ''],
    ['    v', ''],
    ['Layer 2: Decomposition Engine', 'buildTree() -- Differentiated decomposition strategy based on Six-Book type'],
    ['    |', ''],
    ['    v', ''],
    ['Layer 3: Component Annotation', 'annotateTypes() -- Label semantic/phonetic nodes + 3-color phonetic rating + ghost component detection'],
    ['    |', ''],
    ['    v', ''],
    ['Layer 4: Visualization', 'DecompositionGraph component -- D3.js tree layout, color + shape encoding'],
]
make_table(doc,
    ['Architecture Layer', 'Description'],
    flow1_data,
    col_widths=[5, 11],
    header_color='C23B2A')

# --- 1.2 Six-Book Classification Table ---
add_title(doc, '1.2 Six-Book Type x Decomposition Strategy', level=2)

make_table(doc,
    ['Six-Book Type', 'English', 'Decomposition Strategy', 'buildTree() Logic', 'Node Shape'],
    [
        ['Pictographic', 'Xiangxing', 'Atomic node, no split', 'ety.type === pictographic -> return single node', 'Core red circle'],
        ['Indicative', 'Zhishi', 'Atomic node, no split', 'ety.type === indicative -> return single node', 'Core red circle'],
        ['Pictophonetic', 'Xingsheng', 'Single-layer: semantic + phonetic', 'Create 2 child nodes (semantic + phonetic)', 'Semantic: blue circle\nPhonetic: orange diamond'],
        ['Ideographic', 'Huiyi', 'Extract CJK components from IDS', 'Scan CJK chars in decomposition string', 'Gold circle'],
        ['Mutual Explanation', 'Zhuanzhu', 'Not directly processed', 'Covered indirectly via network: same radical + near-homophone + similar meaning', '--'],
        ['Phonetic Loan', 'Jiajie', 'Not directly processed', 'Only shown in detail page via Shuowen data', '--'],
    ],
    col_widths=[2.5, 2.5, 3.5, 5.5, 3],
    header_color='C23B2A')

# --- 1.3 Decomposition Pipeline ---
add_title(doc, '1.3 Decomposition Engine Core Pipeline (for PPT Basic Process SmartArt)', level=2)
add_para(doc, '[PPT Tip: Use "Basic Process" SmartArt, 6 steps in sequence]', size=Pt(9), bold=True)

make_table(doc,
    ['Step', 'Function/Module', 'Input', 'Processing Logic', 'Output'],
    [
        ['1. Receive query', 'decomposeCharacter()', 'Target character', 'Call buildTree() to construct decomposition tree', 'DecompositionTree'],
        ['2. Lookup Six-Book', 'buildTree()', 'Character', 'Read ety.type field from hanzi-dict.json', 'Six-Book type label'],
        ['3. Branch decompose', 'buildTree()', 'Six-Book type + IDS', 'Pictographic/Indicative -> atom;\nPictophonetic -> semantic + phonetic;\nIdeographic -> CJK component list', 'Child node list'],
        ['4. IDS cleaning', 'extractCJK()', 'Raw IDS string', 'Filter non-CJK chars (U+4E00-U+9FFF)\nFilter pure strokes (STROKE_BLACKLIST)', 'Clean component list'],
        ['5. Variant radical annot.', 'getAnnotation()', 'Each component', 'Match 17 variant radical mappings\nSpecial: moon/meat ambiguity, fu/yi ear radical', 'Original radical form'],
        ['6. Render output', 'DecompositionGraph', 'DecompositionTree', 'D3 tree layout -> SVG\nColor code: Red(core), Blue(semantic), Orange(phonetic), Gold(ideographic)', 'Interactive SVG'],
    ],
    col_widths=[2.2, 3.5, 2.5, 5.5, 3],
    header_color='C23B2A')

# --- 1.4 Variant Radical Annotations ---
add_title(doc, '1.4 Variant Radical Annotation System', level=2)

make_table(doc,
    ['Variant Radical', 'Original Radical', 'Description', 'Approx. Coverage'],
    [
        ['Shu Xin Pang (忄)', 'Heart (心)', 'Vertical heart -> heart bottom', '~200+'],
        ['Ti Shou Pang (扌)', 'Hand (手)', 'Lifting hand -> hand radical', '~300+'],
        ['San Dian Shui (氵)', 'Water (水)', 'Three drops -> water radical', '~400+'],
        ['Fan Quan Pang (犭)', 'Dog (犬)', 'Opposing dog -> dog radical', '~80+'],
        ['Si Dian Di (灬)', 'Fire (火)', 'Four dots bottom -> fire bottom', '~50+'],
        ['Dan Ren Pang (亻)', 'Person (人)', 'Single person -> person radical', '~300+'],
        ['Zuo Er Pang (阝L)', 'Fu (阜)', 'Left ear -> fu (mound/hill)', '~100+'],
        ['You Er Pang (阝R)', 'Yi (邑)', 'Right ear -> yi (city)', '~80+'],
        ['Moon/Meat (月)', 'Meat/Moon (肉/月)', 'Ambiguity detection via keyword matching', '~150+'],
        ['Jin (钅)', 'Metal (金)', 'Simplified radical variant', '~200+'],
        ['Shi (饣)', 'Food (食)', 'Simplified radical variant', '~50+'],
        ['Si (纟)', 'Silk (糸)', 'Simplified radical variant', '~100+'],
    ],
    col_widths=[3.5, 3, 6, 3],
    header_color='2D5F8A')

# --- 1.5 Ghost Component System ---
add_title(doc, '1.5 Ghost Component Identification System', level=2)
add_para(doc, '[PPT Tip: Use "Segmented Process" SmartArt to show ghost component detection flow]', size=Pt(9), bold=True)

make_table(doc,
    ['Step', 'Check', 'Example'],
    [
        ['1. Table lookup', 'Search 15 manually annotated high-frequency cases', 'Mai (买) <- cursive simplification of Mai (買), no character-formation meaning'],
        ['2. Annotation', 'getGhostAnnotation() returns ghost description', '"This stroke/component is a simplified derivative form..."'],
        ['3. Rendering', 'DecompositionGraph detects isGhost flag', 'Gray dashed border (#B0ADA5) + fixed tooltip'],
    ],
    col_widths=[3, 6, 7],
    header_color='A39E93')

add_para(doc, '15 ghost component characters: 买 尽 专 长 书 为 东 乐 头 发 后 里 农 龙 万', bold=True, size=Pt(9))

# --- 1.6 Phonetic Reliability Rating ---
add_title(doc, '1.6 Phonetic Reliability Three-Color Rating System', level=2)

make_table(doc,
    ['Rating', 'Condition', 'Initial Match', 'Final Match', 'Icon/Color', 'Example'],
    [
        ['Accurate', 'Initial AND Final both match', 'Yes', 'Yes', 'Green #2E7D32\nLabel: Sheng-V', 'Mu(沐 mu) <- Mu(木 mu)\nInitial+Final identical'],
        ['Approximate', 'Initial OR Final matches', 'Either', 'Either', 'Yellow #F57F17\nLabel: Sheng-~', 'Jiang(江 jiang) <- Gong(工 gong)\nFinal similar'],
        ['Failed', 'Neither initial nor final matches', 'No', 'No', 'Red #C62828\nLabel: Sheng-X', 'He(河 he) <- Ke(可 ke)\nComplete mismatch'],
    ],
    col_widths=[2, 4, 2, 2, 3.5, 4],
    header_color='2D5F8A')

add_para(doc, 'Pinyin parser highlights: recognizes all 21 initials (zh/ch/sh prioritized); zero-initial (an -> initial="", final="an"); u-umlaut handled as v', size=Pt(9))

doc.add_page_break()

# ================================================================
# Part 2: Multi-Dimensional Chinese Character Connection Network
# ================================================================
add_title(doc, 'Part 2: Multi-Dimensional Chinese Character Connection Network', level=1)

# --- 2.1 Overall Architecture ---
add_title(doc, '2.1 System Architecture Flow (for PPT Basic Process SmartArt)', level=2)
add_para(doc, '[PPT Tip: Use "Basic Process" SmartArt with 4 major stages]', size=Pt(9), bold=True)

make_table(doc,
    ['Stage', 'Module/Function', 'Core Task', 'Output'],
    [
        ['Stage 1\nIndex Building', 'loadData()\n(lines 151-262)', 'Build 7 in-memory inverted indices:\nphoneticIndex / semanticIndex / pinyinIndex\npinyinNoToneIndex / reverseIndex\nradicalIndex / structuralRank', '7 Maps of <key, Set<chars>>\nO(1) lookup complexity'],
        ['Stage 2\nRelation Computing', 'computeRelations()\n(lines 953-1084)', 'Real-time compute 9 relation types\n4-level priority sorting + mutual-exclusion dedup\nSimplified-Traditional smart dedup (3-channel)', 'CharRelations object\nTop-N lists per relation type'],
        ['Stage 3\n3D Scoring', 'scoreRelations()\n(lines 1108-1348)', 'Form x Sound x Meaning 3D scoring\nComposite formula: geometric_mean x synergy x freqMod', 'Candidate list sorted by\ntotalScore descending'],
        ['Stage 4\nVisualization', 'CognateGraph component\n(575 lines)', 'D3 force-directed graph layout\n9-color edge encoding + dashed antonyms\nNode size proportional to relation strength', 'Interactive SVG force graph'],
    ],
    col_widths=[2.5, 3.5, 6.5, 4],
    header_color='CA6702')

# --- 2.2 Seven In-Memory Indices ---
add_title(doc, '2.2 Seven In-Memory Inverted Indices', level=2)

make_table(doc,
    ['Index Name', 'Key Type', 'Value Type', 'Purpose', 'Typical Query'],
    [
        ['phoneticIndex', 'Phonetic component', 'Set<Char>', 'Find same-phonetic family', 'phoneticIndex.get("木") -> {沐,霖,林,...}'],
        ['semanticIndex', 'Semantic component', 'Set<Char>', 'Find same-semantic family', 'semanticIndex.get("氵") -> {江,河,海,...}'],
        ['pinyinIndex', 'Pinyin (with tone)', 'Set<Char>', 'Find exact homophones', 'pinyinIndex.get("mu4") -> {木,沐,牧,...}'],
        ['pinyinNoToneIndex', 'Pinyin (no tone)', 'Set<Char>', 'Find near-homophones', 'pinyinNoToneIndex.get("mu") -> {母,木,目,...}'],
        ['reverseIndex', 'Component', 'string[]', 'Find chars containing component', 'reverseIndex.get("木") -> [沐,霖,李,...]'],
        ['radicalIndex', 'Radical', 'string[]', 'Find same-radical chars', 'radicalIndex.get("木") -> [林,森,板,...]'],
        ['structuralRank', 'Component', 'number', 'Measure structural importance', 'structuralRank.get("木") -> high ref count'],
    ],
    col_widths=[3, 2.8, 2.2, 4, 4.5],
    header_color='CA6702')

# --- 2.3 Nine Relation Types ---
add_title(doc, '2.3 Nine Relation Types x Four Priority Levels Matrix', level=2)
add_para(doc, '[Core table -- recommend direct PPT use; can split into "Priority 1-2" and "Priority 3-4" on two slides]', size=Pt(9), bold=True)

make_table(doc,
    ['Priority', 'Relation Type', 'English', 'Computation Logic', 'Edge Color', 'Cap'],
    [
        ['P1\n(Highest)', 'Source Differentiation', 'Differentiation', 'B.phonetic === A AND B.decomposition contains A', 'Red #C23B2A', '30'],
        ['P1\n(Highest)', 'Antonym Pair', 'Antonym', 'Lookup from 57 hand-curated pairs', 'Dark Red #9B2226\n(dashed)', '30'],
        ['P2', 'Same Phonetic Family', 'Phonetic Family', 'Share same phonetic component\n(phoneticIndex lookup)', 'Orange #CA6702', '50'],
        ['P2', 'Same Semantic Family', 'Semantic Family', 'Share same semantic component\n(semanticIndex lookup)', 'Blue #2D5F8A', '50'],
        ['P2', 'Shared Components', 'Shared Components', 'Any CJK component overlap\n(full comparison)', 'Gray #A39E93', '50'],
        ['P3', 'Component Containment', 'Contained In', 'This char appears as component\nin other chars', 'Green #6B7F5E', '40'],
        ['P3', 'Homophones', 'Homophones', 'Initial+Final+Tone all identical\n(pinyinIndex lookup)', 'Gold #8B6914', '20'],
        ['P3', 'Near-Homophones', 'Near-Homophones', 'Initial+Final same, tone different\n(pinyinNoToneIndex lookup)', 'Gold #8B6914', '20'],
        ['P4\n(Lowest)', 'Same Radical Family', 'Radical Family', 'Share same radical\n(radicalIndex lookup)', '--', '40'],
    ],
    col_widths=[1.5, 2.8, 2.5, 4.5, 2.8, 1],
    header_color='CA6702')

# --- 2.4 Mutual Exclusion Dedup ---
add_title(doc, '2.4 Mutual Exclusion Dedup Mechanism (takeList Algorithm)', level=2)
add_para(doc, '[PPT Tip: Use "Process" SmartArt to show 5-step dedup flow]', size=Pt(9), bold=True)

make_table(doc,
    ['Step', 'Operation', 'Description'],
    [
        ['1. Initialize', 'claimed = new Set<string>()', 'Create empty set to track chars already claimed by higher priority'],
        ['2. Iterate by priority', 'for each priority level (P1 -> P4)', 'Process from high to low, ensuring high-priority relations shown first'],
        ['3. Sort by importance', 'Sort candidates by structuralRank descending', 'More fundamental, frequently-used chars appear first'],
        ['4. Simp-Trad dedup', 'dedupSimpTrad() filter', 'Three-channel strategy to remove simplified-traditional variants (see 2.5)'],
        ['5. Take Top-N', 'Remove claimed -> Take Top-N -> Add to claimed', 'Each char appears only once, in its highest-priority relation'],
    ],
    col_widths=[2.5, 5.5, 8],
    header_color='6B7F5E')

# --- 2.5 Simp-Trad Dedup ---
add_title(doc, '2.5 Simplified-Traditional Smart Dedup: Three-Channel Strategy', level=2)

make_table(doc,
    ['Channel', 'Strategy', 'Implementation', 'Example'],
    [
        ['Channel 1\nExplicit Mapping', 'Lookup simp-trad-map.json', 'Direct search for known simplified-traditional pairs', 'Guo(国 simp) <-> Guo(國 trad)'],
        ['Channel 2\nPinyin+Def Inference', 'Same pinyin (no tone) + identical definition', 'Inferred as variant relationship', 'Wei(为 simp) <-> Wei(爲 trad)\n(pinyin wei, same definition)'],
        ['Channel 3\nCascade Propagation', 'Radical trad-mark -> Component trad-mark\n-> Secondary propagation', 'Marked chars propagate to components', 'e.g. Yan(言) -> Jianyan(讠)\npropagates to all chars with 言'],
    ],
    col_widths=[2.5, 4.5, 5, 4],
    header_color='6B7F5E')

# --- 2.6 Three-Dimensional Scoring ---
add_title(doc, '2.6 Form-Sound-Meaning 3D Affinity Scoring Algorithm', level=2)
add_para(doc, '[Core algorithm -- PPT Tip: Use "Process" SmartArt with 3 parallel dimensions]', size=Pt(9), bold=True)

make_table(doc,
    ['Dimension', 'Max Score', 'Scoring Rules (partial)', 'Role in Formula'],
    [
        ['Form\n(0-100)', '100',
         '- Shared component Jaccard: sharedCount/max(a,b)*50 (max 50)\n- Same radical: +25\n- Same phonetic: +20\n- Differentiation relation: +15\n- Component containment (bidirectional): +10',
         'Geometric mean component\nFloor value = 15'],
        ['Sound\n(0-100)', '100',
         '- Exact homophone (initial+final+tone all same): 100\n- Near-homophone (same syllable, different tone): 50\n- No relation: 0',
         'Geometric mean component\nFloor value = 15'],
        ['Meaning\n(0-100)', '100',
         '- Same semantic component: +40\n- Antonym pair: +40\n- Source differentiation: +30\n- Same radical: +15\n- Component containment: +10',
         'Geometric mean component\nFloor value = 15'],
    ],
    col_widths=[2.5, 1.8, 8, 4],
    header_color='8B6914')

add_para(doc, '', size=Pt(6))
add_para(doc, 'Composite Scoring Formula', bold=True, size=Pt(10))

make_table(doc,
    ['Formula Component', 'Expression', 'Explanation'],
    [
        ['Geometric Mean', 'geoMean = cbrt(f * s * m)\nf=max(15,formScore), s=max(15,soundScore), m=max(15,meaningScore)',
         'Naturally penalizes single-dimension matching\nExample: cbrt(90*15*15)=27 vs cbrt(50*50*30)=42'],
        ['Synergy Bonus', 'dims = (form>20?1:0)+(sound>20?1:0)+(meaning>20?1:0)\nsynergy = dims>=3 ? 1.5 : dims>=2 ? 1.2 : 1.0',
         'Rewards multi-dimensional correlation\n3 dims >20 -> 1.5x\n2 dims >20 -> 1.2x'],
        ['Frequency Modifier', 'importance = structuralRank + hasEtymology?5:0 + log2(radicalFamilySize+1)\nfreqMod = min(1.35, 0.35 + log2(importance+1)/6)',
         'Boosts fundamental common chars\nRange 0.35~1.35'],
        ['Final Score', 'totalScore = round(geometricMean * synergy * freqMod)', 'All dimensions combine synergistically\nOutput sorted descending by this value'],
    ],
    col_widths=[2.8, 8, 5.5],
    header_color='8B6914')

# --- 2.7 End-to-End Data Flow ---
add_title(doc, '2.7 End-to-End Data Flow: From Input to Visualization', level=2)
add_para(doc, '[PPT Tip: Use "Basic Process" SmartArt to show complete data pipeline]', size=Pt(9), bold=True)

make_table(doc,
    ['Sequence', 'Module', 'Operation', 'Data Volume / Performance'],
    [
        ['1. User Input', 'Explore Page', 'User enters character (e.g. 沐) or English word', '--'],
        ['2. Data Loading', 'loadData()', 'fetch hanzi-dict.json + hanzi-index.json\nBuild 7 in-memory indices', 'Core data ~500KB (gzip)\n3G network ~3-5s'],
        ['3. Decomp Query', 'decomposeCharacter()', 'Lookup Six-Book type -> buildTree() -> annotateTypes()', '<10ms (in-memory lookup)'],
        ['4. Render Decomp', 'DecompositionGraph', 'D3 tree layout -> SVG output\nColor/shape/label encoding', 'Single char render <50ms'],
        ['5. Click Phonetic', 'User Interaction', 'Click phonetic node of pictophonetic char -> triggers connection query', '--'],
        ['6. Relation Compute', 'computeRelations()', 'Real-time compute 9 relations\n4-level priority + mutual dedup + simp-trad dedup', '<50ms (index lookup)'],
        ['7. 3D Scoring', 'scoreRelations()', 'Score each candidate\ngeometric_mean x synergy x freqMod', '<30ms'],
        ['8. Render Network', 'CognateGraph', 'D3 force graph -> 9-color edges + nodes\nSimulation auto-stops after 3s', '~40 node cap\nRender <200ms'],
    ],
    col_widths=[1.8, 3.2, 6, 4.5],
    header_color='C23B2A')

# --- 2.8 Build-Time vs Runtime ---
add_title(doc, '2.8 Build-Time vs Runtime Data Flow', level=2)

make_table(doc,
    ['Phase', 'Trigger', 'Execution', 'Output Artifact', 'Used at Runtime?'],
    [
        ['Build-Time', 'npm run build\n(Vite build)', 'build-relations.ts\nCompute all char relations\nBuild indices', 'char-relations.json\n(static JSON)', 'No (currently real-time)\nPre-computed file as backup'],
        ['Runtime\n(Page Load)', 'First page load', 'loadData()\nfetch JSON + build indices', '7 in-memory Maps', 'Yes'],
        ['Runtime\n(User Query)', 'User search/click', 'computeRelations()\nscoreRelations()\nReal-time compute + sort', 'CharRelations object\n-> D3 visualization', 'Yes'],
    ],
    col_widths=[2.5, 3, 4, 3.5, 4],
    header_color='2D5F8A')

doc.add_page_break()

# ================================================================
# Part 3: Relationship Between the Two Systems
# ================================================================
add_title(doc, 'Part 3: Two-System Relationship Overview', level=1)

add_title(doc, '3.1 Positioning of Two Systems in the Project', level=2)

make_table(doc,
    ['Dimension', 'Six-Book Decomposition Engine', 'Multi-Dimensional Connection Network'],
    [
        ['Core Goal', 'Vertical depth: explore internal structure of a single character',
         'Horizontal breadth: discover multi-dimensional relationships between characters'],
        ['Metaphor', 'Scalpel -- dissect character formation rationale',
         'Relational Web -- weave character family connections'],
        ['Core Functions', 'buildTree() + annotateTypes()',
         'computeRelations() + scoreRelations()'],
        ['Visual Component', 'DecompositionGraph (tree layout)',
         'CognateGraph (force-directed graph)'],
        ['D3 Layout', 'Tree layout (hierarchical expansion)',
         'Force-directed graph (free expansion)'],
        ['Node Color Code', 'Red(core)/Blue(semantic)/Orange(phonetic)/Gold(ideographic)',
         '9-color edges + node size proportional to relation strength'],
        ['Special Markers', 'Ghost components (dashed border) + phonetic rating (3-color)',
         'Antonym relations (dashed lines) + simp-trad dedup'],
        ['Data Source', 'hanzi-dict.json (Six-Book classification + IDS)',
         '7 in-memory inverted indices (real-time computation)'],
        ['User Entry Point', 'Click character -> auto-show decomposition',
         'Click phonetic/semantic component -> explore related character network'],
        ['Theoretical Basis', 'Shuowen Jiezi Six-Book theory\nWang Ning character formation theory (component hierarchy)',
         'Wang Ning character formation theory (character family connection)\nShen Jianshi right-side theory'],
    ],
    col_widths=[3, 6.5, 6.5],
    header_color='2D5F8A')

add_title(doc, '3.2 Data Flow Connection Points', level=2)
add_para(doc, '[PPT Tip: Use "Basic Process" SmartArt to show how the two systems connect]', size=Pt(9), bold=True)

make_table(doc,
    ['Connection', 'From Decomposition Engine ->', 'To Connection Network ->', 'User Experience'],
    [
        ['Connection 1\nSemantic Click', 'Decomp graph shows semantic component 氵\nlabeled as blue semantic node',
         'semanticIndex query\nSame semantic family: 江/河/海/湖...',
         'User understands water-radical\ncharacter family semantic relations'],
        ['Connection 2\nPhonetic Click', 'Decomp graph shows phonetic component 木\nlabeled as orange phonetic node + rating',
         'phoneticIndex query\nSame phonetic family: 沐/霖/林...',
         'User discovers mu-sound character\nfamily phonetic patterns'],
        ['Connection 3\nIdeographic Component', 'Decomp graph shows ideographic components\ne.g. 信 -> 人 + 言',
         'reverseIndex query\nComponent containment: 人->从/众; 言->语/说',
         'User explores component reuse\nacross different characters'],
    ],
    col_widths=[2.5, 5.5, 5, 4],
    header_color='C23B2A')

doc.add_page_break()

# ================================================================
# Part 4: Bloom's Taxonomy Mapping
# ================================================================
add_title(doc, 'Part 4: Bloom\'s Taxonomy x Product Feature Mapping (PPT Matrix Material)', level=1)

make_table(doc,
    ['Cognitive Level', 'Bloom\'s Taxonomy', 'Product Feature', 'Interaction Mode', 'Cognitive Output'],
    [
        ['Remembering', 'Remembering', 'Character Detail Page', 'Static information display', 'Recognize glyph, pronunciation, definition'],
        ['Understanding', 'Understanding', 'Component Decomposition Graph', 'Hover expand + click navigate', 'Understand character internal formation rationale'],
        ['Applying', 'Applying', 'Glyph Evolution Timeline', 'Horizontal swipe browse', 'Recognize historical glyph evolution patterns'],
        ['Analyzing', 'Analyzing', 'Cognate Relation Network Graph', 'Drag + zoom + hover highlight', 'Analyze multi-dimensional relations among character family members'],
        ['Evaluating', 'Evaluating', 'Phonetic Reliability 3-Color Rating', 'Click phonetic to view rating', 'Judge phonetic component reliability'],
        ['Creating', 'Creating', 'Character Assembly Game', 'Drag components to assemble/disassemble', 'Apply formation knowledge to create characters'],
    ],
    col_widths=[2.5, 3, 4, 4, 4.5],
    header_color='2D5F8A')

# ================================================================
# Save
# ================================================================
doc.save(OUTPUT_PATH)
print(f'Done! Document saved to: {OUTPUT_PATH}')
print(f'File size: {os.path.getsize(OUTPUT_PATH) / 1024:.1f} KB')
