# -*- coding: utf-8 -*-
"""
Generate editable Word document with native tables, flowcharts, and diagrams.
All content is 100% editable — tables, text, colors can be modified directly in Word.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

OUTPUT_PATH = r'C:\文件\大三\大三下\字里行间\新建文件夹\app\PPT流程图_v4_最终版.docx'

doc = Document()

# ============================================================
# Global Styles
# ============================================================
style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(10)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

for section in doc.sections:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

# ============================================================
# Color Palette
# ============================================================
C_RED    = 'C23B2A'
C_DARKRED = '9B2226'
C_BLUE   = '2D5F8A'
C_ORANGE = 'CA6702'
C_GOLD   = '8B6914'
C_GREEN  = '2E7D32'
C_GRAY   = 'A39E93'
C_LIGHT  = 'F2F0EB'
C_WHITE  = 'FFFFFF'
C_DARK   = '2C2C2C'
C_YELLOW = 'F57F17'
C_CYAN   = '0E7C7B'
C_PURPLE = '6B4E71'

# ============================================================
# Helper Functions
# ============================================================

def set_cell_bg(cell, color):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def hex_to_rgb(hex_str):
    """Convert hex color string to RGBColor, handling both raw hex and RGBColor input."""
    if isinstance(hex_str, RGBColor):
        return hex_str
    h = hex_str.lstrip('#')
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))

def set_cell_text(cell, text, bold=False, size=Pt(9), color=None, align=WD_ALIGN_PARAGRAPH.CENTER, font_name='Microsoft YaHei'):
    """Set cell text with formatting."""
    # Clear existing
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    # Remove space
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = Pt(13)
    # Add run
    run = p.add_run(text)
    run.bold = bold
    run.font.size = size
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if color:
        if isinstance(color, str):
            color = hex_to_rgb(color)
        run.font.color.rgb = color


def add_title(text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')


def add_para(text, bold=False, size=Pt(10)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = size
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    return p


def add_sep():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('—' * 40)
    r.font.size = Pt(7)
    r.font.color.rgb = RGBColor(0xA3, 0x9E, 0x93)


def new_table(rows, cols, col_widths=None):
    """Create a new table with grid style."""
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return table


def header_row(table, texts, bg_color=C_RED, text_color=None):
    """Format first row as header."""
    if text_color is None:
        text_color = RGBColor(0xFF, 0xFF, 0xFF)
    for i, text in enumerate(texts):
        cell = table.rows[0].cells[i]
        set_cell_bg(cell, bg_color)
        set_cell_text(cell, text, bold=True, size=Pt(9), color=text_color)
    return table


def data_rows(table, data, colors=None):
    """Fill data rows with alternating background."""
    for r, row_data in enumerate(data):
        row = table.rows[r + 1]
        for c, text in enumerate(row_data):
            cell = row.cells[c]
            set_cell_text(cell, str(text), size=Pt(8.5))
            if r % 2 == 0:
                set_cell_bg(cell, C_LIGHT)
            if colors and c in colors:
                # Apply cell-specific color if provided
                pass
    return table


def make_info_table(headers, data, col_widths=None, header_color=C_RED):
    """Quickly create a complete info table."""
    table = new_table(len(data) + 1, len(headers), col_widths)
    header_row(table, headers, bg_color=header_color)
    data_rows(table, data)
    doc.add_paragraph()
    return table


def flow_block(text_lines, color, width_cm=3.5):
    """Create a flow block as a single-cell colored table."""
    t = doc.add_table(rows=1, cols=1)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.rows[0].cells[0]
    cell.width = Cm(width_cm)
    set_cell_bg(cell, color)
    set_cell_text(cell, text_lines, bold=True, size=Pt(10), color=RGBColor(0xFF,0xFF,0xFF))
    return t


def arrow_down():
    """Add a centered down arrow."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('▼')
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0xA3, 0x9E, 0x93)
    r.font.name = 'Microsoft YaHei'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')


def arrow_right():
    """Add a centered right arrow."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('▶')
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0xA3, 0x9E, 0x93)
    r.font.name = 'Microsoft YaHei'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')


def spacer(h=Pt(4)):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(' ')
    r.font.size = h


# ================================================================
# COVER PAGE
# ================================================================
add_title('LINES 字里行间', level=0)
add_title('PPT 流程图与数据图表（可编辑版）', level=1)
add_para('本文件所有表格、流程图、文字均为 Word 原生格式，可直接编辑修改。', bold=True, size=Pt(11))
add_para('使用方法：选中任意表格 → 修改文字/颜色/布局 → 复制粘贴到 PPT', size=Pt(10))
add_para('生成日期：2026-06-03', size=Pt(9))
doc.add_page_break()


# ================================================================
# PART 1: 拆解引擎
# ================================================================
add_title('第一部分：六书驱动的汉字智能拆解引擎', level=1)

# === 1.1 四层架构流程图 ===
add_title('1.1 系统四层架构（可编辑流程图）', level=2)
add_para('提示：每个色块是一个独立的 1×1 表格，可直接点击修改文字。▼ 箭头可删除或替换。', size=Pt(8))

flow_block('第一层：数据源\n\nhanzi-dict.json\n六书分类 + IDS + 说文解字', C_RED, 10)
arrow_down()
flow_block('第二层：拆解引擎\n\nbuildTree()\n根据六书类型执行差异化拆解', C_ORANGE, 10)
arrow_down()
flow_block('第三层：部件标注\n\nannotateTypes()\n语义/语音标注 + 幽灵检测 + 声旁评级', C_BLUE, 10)
arrow_down()
flow_block('第四层：可视化渲染\n\nDecompositionGraph 组件\nD3 树形图 · 颜色编码 · 交互导航', C_GREEN, 10)

spacer()
doc.add_page_break()

# === 1.2 六书拆解策略表 ===
add_title('1.2 六书分类 × 拆解策略对照表', level=2)
make_info_table(
    ['六书类型', '英文名', '拆解策略', 'buildTree() 逻辑', '节点形态', '颜色'],
    [
        ['象形字', 'Pictographic', '原子节点\n不拆分', 'type === pictographic\n→ 直接返回单节点', '● 圆形', '红色'],
        ['指事字', 'Indicative', '原子节点\n不拆分', 'type === indicative\n→ 直接返回单节点', '● 圆形', '红色'],
        ['形声字', 'Pictophonetic', '单层展开\n形旁 + 声旁', '创建两个子节点\n语义 + 语音', '● 语义圆形\n◆ 语音菱形', '蓝 + 橙'],
        ['会意字', 'Ideographic', 'IDS 提取\nCJK 部件', '扫描 decomposition\n中的 CJK 字符', '● 圆形', '金色'],
        ['转注字', 'Mutual Expl.', '间接覆盖\n系联网络', '不直接处理\n网络中间接体现', '—', '—'],
        ['假借字', 'Phonetic Loan', '仅展示\n详情页', '不纳入拆解引擎\n详情页展示', '—', '—'],
    ],
    col_widths=[2.0, 2.2, 2.5, 4.0, 2.5, 1.5],
    header_color=C_RED
)

# === 1.3 拆解流水线 ===
add_title('1.3 拆解引擎处理流水线（可编辑流程图）', level=2)
add_para('提示：6 个色块各自是独立的 1x1 表格，可分别编辑。▶ 箭头为文本符号。', size=Pt(8))

# Build a horizontal flow using a single-row table
pipeline_table = new_table(3, 6, col_widths=[2.0, 2.0, 2.0, 2.0, 2.0, 2.0])
steps = [
    ('① 输入汉字', C_RED),
    ('② 查六书类型', C_ORANGE),
    ('③ 分支拆解', C_BLUE),
    ('④ IDS清洗', C_GOLD),
    ('⑤ 变形标注\n  +幽灵检测', C_GREEN),
    ('⑥ D3渲染', C_PURPLE),
]
for i, (label, color) in enumerate(steps):
    cell = pipeline_table.rows[1].cells[i]
    set_cell_bg(cell, color)
    set_cell_text(cell, label, bold=True, size=Pt(10), color=RGBColor(0xFF,0xFF,0xFF))
    # Add arrow in row 0
    if i > 0:
        set_cell_text(pipeline_table.rows[0].cells[i], '▶', bold=True, size=Pt(12), color=RGBColor(0xA3,0x9E,0x93))
    # Add function name in row 2
    func_names = ['decompose()', 'buildTree()', 'buildTree()', 'extractCJK()', 'getAnnotation()', 'DecompGraph']
    set_cell_text(pipeline_table.rows[2].cells[i], func_names[i], size=Pt(7.5), color=RGBColor(0x88,0x88,0x88))

doc.add_paragraph()
doc.add_page_break()

# === 1.4 声旁三色评级 ===
add_title('1.4 声旁可靠性三色评级系统', level=2)

make_info_table(
    ['评级', '图标', '判断条件', '声母', '韵母', '颜色', '示例'],
    [
        ['准确\nAccurate', '声✓', '声母 AND 韵母\n均匹配', '✓ 匹配', '✓ 匹配', '🟢 绿色\n#2E7D32', '沐 mù ← 木 mù\n声韵全同 → 准确'],
        ['近似\nApprox.', '声~', '声母 OR 韵母\n其一匹配', '可匹配\n或不匹配', '可匹配\n或不匹配', '🟡 黄色\n#F57F17', '江 jiāng ← 工 gōng\n韵母相近 → 近似'],
        ['失效\nFailed', '声✗', '声韵均不匹配', '✗ 不匹配', '✗ 不匹配', '🔴 红色\n#C62828', '河 hé ← 可 kě\n声韵全异 → 失效'],
    ],
    col_widths=[1.8, 1.2, 3.0, 1.8, 1.8, 2.0, 4.0],
    header_color=C_BLUE
)

# === 1.5 变形偏旁 ===
add_title('1.5 变形偏旁标注系统', level=2)

make_info_table(
    ['变形偏旁', '原始偏旁', '说明', '特殊处理'],
    [
        ['忄 → 心', '心', '竖心旁 → 心字底', '—'],
        ['扌 → 手', '手', '提手旁 → 手字旁', '—'],
        ['氵 → 水', '水', '三点水 → 水字旁', '—'],
        ['犭 → 犬', '犬', '反犬旁 → 犬字旁', '—'],
        ['灬 → 火', '火', '四点底 → 火字底', '—'],
        ['亻 → 人', '人', '单人旁 → 人字旁', '—'],
        ['阝(左) → 阜', '阜', '左耳旁 → 阜(山丘)', '左右位置判定'],
        ['阝(右) → 邑', '邑', '右耳旁 → 邑(城邑)', '左右位置判定'],
        ['月 → 肉/月', '肉/月', '月字旁二义性', '关键词匹配\n身体词汇 → 肉'],
        ['钅 → 金', '金', '简化部首', '—'],
        ['饣 → 食', '食', '简化部首', '—'],
        ['纟 → 糸', '糸', '简化部首', '—'],
    ],
    col_widths=[3.0, 2.0, 5.0, 4.0],
    header_color=C_BLUE
)

doc.add_page_break()

# === 1.6 幽灵部件 ===
add_title('1.6 幽灵部件识别系统', level=2)

make_info_table(
    ['判定步骤', '函数/模块', '说明', '示例'],
    [
        ['① 查表匹配', 'getGhostAnnotation()', '在 15 个人工标注的\n高频用例中查找', '买 → 草书楷化（買→买）'],
        ['② 标注描述', 'ghostComponents.ts', '返回幽灵描述信息', '"该笔画/部件为\n简体简化衍生形态"'],
        ['③ 渲染标记', 'DecompositionGraph', '检测 isGhost 标志\n→ 灰色虚线边框', '虚线边框 (#B0ADA5)\n+ 固定提示语'],
    ],
    col_widths=[2.5, 4.0, 4.5, 4.5],
    header_color=C_GRAY
)

add_para('幽灵部件覆盖的 15 个高频字：买、尽、专、长、书、为、东、乐、头、发、后、里、农、龙、万', bold=True, size=Pt(9))

doc.add_page_break()

# ================================================================
# PART 2: 系联网络（精简版，适配 PPT）
# ================================================================
add_title('第二部分：多维汉字系联网络', level=1)

# === 2.1 四阶段架构 ===
add_title('2.1 系统四阶段架构', level=2)

flow_block('阶段一：索引构建\nloadData()\n7 个倒排索引', C_BLUE, 7)
arrow_down()
flow_block('阶段二：关系计算\ncomputeRelations()\n9 种关系 · 4 级优先级', C_ORANGE, 7)
arrow_down()
flow_block('阶段三：三维评分\nscoreRelations()\n字形 × 字音 × 字义', C_GOLD, 7)
arrow_down()
flow_block('阶段四：可视化\nCognateGraph\nD3 力导向图', C_GREEN, 7)

spacer()

# === 2.2 七大索引（精简） ===
add_title('2.2 七大倒排索引', level=2)

make_info_table(
    ['索引', '用途', '示例'],
    [
        ['phoneticIndex', '同声旁字族', '木 → {沐, 霖, 林}'],
        ['semanticIndex', '同形旁字族', '氵 → {江, 河, 海}'],
        ['pinyinIndex', '同音字', 'mu4 → {木, 沐, 牧}'],
        ['pinyinNoToneIdx', '近音字', 'mu → {母, 木, 目}'],
        ['reverseIndex', '包含某部件的字', '木 → [沐, 霖, 李]'],
        ['radicalIndex', '同部首字', '木部 → [林, 森, 板]'],
        ['structuralRank', '部件重要度', '木 → 高引用计数'],
    ],
    col_widths=[3.5, 5.0, 6.0],
    header_color=C_ORANGE
)

# === 2.3 九种关系类型（核心大表） ===
add_title('2.3 九种关系 × 四级优先级', level=2)

make_info_table(
    ['优先级', '关系类型', '计算逻辑', '边颜色', '上限'],
    [
        ['P1', '源流分化', 'B.声旁=A 且 B含A', '红 #C23B2A', '30'],
        ['P1', '反义对', '57对手工配对查表', '深红 #9B2226', '30'],
        ['P2', '同声旁族', '共享表音构件', '橙 #CA6702', '50'],
        ['P2', '同形旁族', '共享表意构件', '蓝 #2D5F8A', '50'],
        ['P2', '共享构件', 'CJK部件重叠', '灰 #A39E93', '50'],
        ['P3', '构件包含', '该字为他字部件', '绿 #6B7F5E', '40'],
        ['P3', '同音字', '声韵调全同', '金 #8B6914', '20'],
        ['P3', '近音字', '声韵同·调不同', '金虚线', '20'],
        ['P4', '同部首族', '共享部首', '—', '40'],
    ],
    col_widths=[1.5, 2.8, 5.0, 3.0, 1.5],
    header_color=C_ORANGE
)

spacer()

# === 2.4 互斥去重 ===
add_title('2.4 互斥去重机制', level=2)

dedup_table = new_table(1, 5, col_widths=[2.6, 2.6, 2.6, 2.6, 2.6])
dedup_steps = [
    ('① 初始化\nclaimed=Set()', C_DARK),
    ('② 按优先级\nP1→P4', C_RED),
    ('③ 重要度排序\nstructuralRank', C_ORANGE),
    ('④ 简繁去重\ndedupSimpTrad()', C_BLUE),
    ('⑤ 取Top-N\n加入claimed', C_GREEN),
]
for i, (label, color) in enumerate(dedup_steps):
    cell = dedup_table.rows[0].cells[i]
    set_cell_bg(cell, color)
    set_cell_text(cell, label, bold=True, size=Pt(9), color=RGBColor(0xFF,0xFF,0xFF))
doc.add_paragraph()

add_para('每个汉字仅在其最高优先级关系中亮相一次。', bold=True, size=Pt(9))

# === 2.5 简繁去重 ===
add_title('2.5 简繁智能去重', level=2)

simp_table = new_table(1, 3, col_widths=[4.5, 4.5, 4.5])
channels = [
    ('通道一：显式映射\n查 simp-trad-map.json\n例：国 ↔ 國', C_BLUE),
    ('通道二：拼音+释义推断\n同拼音+同释义 → 异体\n例：为 ↔ 爲', C_ORANGE),
    ('通道三：标记级联传播\n部首标记 → 部件标记 → 全局\n例：言 → 讠', C_GREEN),
]
for i, (label, color) in enumerate(channels):
    cell = simp_table.rows[0].cells[i]
    set_cell_bg(cell, color)
    set_cell_text(cell, label, bold=False, size=Pt(9), color=RGBColor(0xFF,0xFF,0xFF))
doc.add_paragraph()

# === 2.6 三维评分（PPT 信息图风格） ===
add_title('2.6 形音义三维亲密度评分', level=2)

# ---- Row 1: Three dimension cards (ultra-compact) ----
dim_table = new_table(1, 3, col_widths=[4.5, 4.5, 4.5])

dim_data = [
    ('字形 Form', '共享构件 · 同部首 · 同声旁\n分化关系 · 构件包含', C_BLUE),
    ('字音 Sound', '完全同音 = 100\n近音(同音节不同调) = 50', C_ORANGE),
    ('字义 Meaning', '同形旁 · 反义对\n分化 · 同部首 · 构件包含', C_GOLD),
]
for i, (title, rules, color) in enumerate(dim_data):
    cell = dim_table.rows[0].cells[i]
    set_cell_bg(cell, color)
    set_cell_text(cell, f'{title}\n\n{rules}', bold=False, size=Pt(9), color=RGBColor(0xFF,0xFF,0xFF))

spacer(Pt(6))

# ---- Row 2: Core formula pipeline (the key visual) ----
add_para('评分公式', bold=True, size=Pt(10))

formula = new_table(1, 5, col_widths=[2.5, 1.0, 2.5, 1.0, 2.5])
f_steps = [
    ('∛(f × s × m)\n几何平均', C_BLUE),
    ('×', None),
    ('1.0 / 1.2 / 1.5\n协同加成', C_ORANGE),
    ('=', None),
    ('Total Score\n综合亲密度', C_RED),
]
for i, (text, color) in enumerate(f_steps):
    cell = formula.rows[0].cells[i]
    if color:
        set_cell_bg(cell, color)
        set_cell_text(cell, text, bold=True, size=Pt(9), color=RGBColor(0xFF,0xFF,0xFF))
    else:
        set_cell_text(cell, text, bold=True, size=Pt(16), color=RGBColor(0x66,0x66,0x66))

spacer(Pt(6))

# ---- Row 3: Key insight + Example (side by side) ----
insight_table = new_table(1, 2, col_widths=[7.0, 7.0])

cell = insight_table.rows[0].cells[0]
set_cell_bg(cell, C_LIGHT)
set_cell_text(cell,
    '设计思想\n\n'
    '几何平均 ∛(f·s·m) 自然惩罚「偏科」\n'
    '  例：∛(90×15×15) = 27  (单维高分，总分低)\n'
    '  例：∛(50×50×30) = 42  (三维均衡，总分高)\n\n'
    '协同加成奖励「多维印证」\n'
    '  3 个维度均 >20 → ×1.5\n'
    '  2 个维度均 >20 → ×1.2\n\n'
    '频率修正提升基础字权重 (0.35~1.35)',
    bold=False, size=Pt(8.5), color=C_DARK, align=WD_ALIGN_PARAGRAPH.LEFT)

cell = insight_table.rows[0].cells[1]
set_cell_bg(cell, C_DARK)
set_cell_text(cell,
    '计算实例：沐 ← 木\n\n'
    '字形 f = 90   (共享构件 + 同声旁)\n'
    '字音 s = 100  (完全同音 mù = mù)\n'
    '字义 m = 60   (同形旁 氵 → 水)\n'
    '━━━━━━━━━━━━━━━━\n'
    '∛(90×100×60) = 81.5\n'
    '       × 1.5  (三维 >20)\n'
    '       = 122  综合得分',
    bold=False, size=Pt(8.5), color=RGBColor(0xFF,0xFF,0xFF), align=WD_ALIGN_PARAGRAPH.LEFT)

spacer(Pt(4))

# ---- Row 4: Score spectrum ----
add_para('得分区间', bold=True, size=Pt(10))

spectrum = new_table(1, 5, col_widths=[3.0, 2.5, 3.0, 2.5, 3.0])
spec_data = [
    ('0 ~ 30\n弱关联', C_GRAY),
    ('30 ~ 60\n中关联', C_GOLD),
    ('60 ~ 90\n强关联', C_ORANGE),
    ('90 ~ 120\n紧密关联', C_RED),
    ('120+\n核心字族', C_DARKRED),
]
for i, (text, color) in enumerate(spec_data):
    cell = spectrum.rows[0].cells[i]
    set_cell_bg(cell, color)
    set_cell_text(cell, text, bold=True, size=Pt(9), color=RGBColor(0xFF,0xFF,0xFF))

spacer()

# === 2.7 端到端数据流 ===
add_title('2.7 端到端数据流', level=2)

pipeline2 = new_table(1, 8, col_widths=[1.7, 1.7, 1.7, 1.7, 1.7, 1.7, 1.7, 1.7])
flow_steps = [
    ('输入\n汉字', C_DARK),
    ('加载\n~500KB', C_BLUE),
    ('拆解\n<10ms', C_RED),
    ('D3\n树形图', C_GREEN),
    ('点击\n声旁', C_DARK),
    ('关系\n计算', C_ORANGE),
    ('三维\n评分', C_GOLD),
    ('D3\n力导向图', C_GREEN),
]
for i, (label, color) in enumerate(flow_steps):
    cell = pipeline2.rows[0].cells[i]
    set_cell_bg(cell, color)
    set_cell_text(cell, label, bold=True, size=Pt(8.5), color=RGBColor(0xFF,0xFF,0xFF))
doc.add_paragraph()

spacer()

# ================================================================
# PART 3: 系统关系总览
# ================================================================
add_title('第三部分：系统关系总览', level=1)

# === 3.1 两大系统对比 ===
add_title('3.1 两大系统定位对比表', level=2)

make_info_table(
    ['对比维度', '六书驱动拆解引擎', '多维汉字系联网络'],
    [
        ['核心目标', '纵向深入：探索单字内部结构', '横向展开：发现字间多维关联'],
        ['核心隐喻', '解剖刀 —— 解析构形理据', '关系网 —— 编织字族脉络'],
        ['核心函数', 'buildTree()\nannotateTypes()', 'computeRelations()\nscoreRelations()'],
        ['可视化组件', 'DecompositionGraph\n(D3 树形图)', 'CognateGraph\n(D3 力导向图)'],
        ['布局方式', '层级树形（父子展开）', '力导向图（自由展开）'],
        ['节点颜色', '红(核心) 蓝(语义)\n橙(语音) 金(会意)', '9 色边编码\n节点大小 ∝ 关系强度'],
        ['特殊标记', '幽灵部件(虚线边框)\n声旁评级(三色标签)', '反义关系(虚线边)\n简繁智能去重'],
        ['数据来源', 'hanzi-dict.json\n(六书分类 + IDS)', '7 个内存倒排索引\n(实时计算)'],
        ['用户入口', '点击汉字\n→ 自动展示拆解', '点击声旁/形旁\n→ 探索关联字网络'],
        ['理论依据', '《说文》六书理论\n王宁汉字构形学', '王宁字族系联\n沈兼士右文说'],
    ],
    col_widths=[3.0, 6.0, 6.0],
    header_color=C_BLUE
)

# === 3.2 数据流衔接 ===
add_title('3.2 系统衔接：从拆解到系联', level=2)

make_info_table(
    ['衔接点', '拆解引擎输出', '系联网络输入', '用户体验'],
    [
        ['形旁点击', '展示语义节点"氵"\n标注为蓝色圆形', '→ semanticIndex 查询\n→ 同形旁族：江/河/海/湖', '理解水部字族\n语义关联规律'],
        ['声旁点击', '展示语音节点"木"\n标注为橙色菱形+评级', '→ phoneticIndex 查询\n→ 同声旁族：沐/霖/林', '发现木声字族\n语音规律'],
        ['会意部件', '展示会意部件\n如「信」→「人」「言」', '→ reverseIndex 查询\n→ 构件包含：人→从/众', '探索部件在不同字\n中的复用规律'],
    ],
    col_widths=[2.5, 5.0, 5.0, 4.0],
    header_color=C_RED
)

doc.add_page_break()

# === 3.3 布鲁姆认知映射 ===
add_title('3.3 布鲁姆认知目标 × 产品功能映射', level=2)
add_para('【可直接复制到 PPT 中作为矩阵图使用】', bold=True, size=Pt(9))

make_info_table(
    ['认知层级', "Bloom's Taxonomy", '产品功能', '交互方式', '认知产出'],
    [
        ['记忆', 'Remembering', '汉字详情页', '静态信息展示', '识记字形·读音·释义'],
        ['理解', 'Understanding', '部件拆解图', '悬停展开+单击导航', '理解内部构形理据'],
        ['应用', 'Applying', '字形演变时间轴', '横向滑动浏览', '认识字形历史演变'],
        ['分析', 'Analyzing', '同源字关系网络图', '拖拽+缩放+悬停高亮', '分析字族多维关系'],
        ['评价', 'Evaluating', '声旁可靠性三色评级', '点击声旁查看评级', '判断声旁表音可靠性'],
        ['创造', 'Creating', '汉字拼拆游戏', '拖拽部件拼拆', '运用构形知识创造'],
    ],
    col_widths=[2.0, 3.0, 4.0, 4.0, 4.0],
    header_color=C_BLUE
)

# === 3.4 技术栈总览 ===
add_title('3.4 技术架构一览', level=2)

make_info_table(
    ['层级', '技术方案', '说明'],
    [
        ['前端框架', 'React 18 + TypeScript', '组件化开发，类型安全'],
        ['可视化', 'D3.js v7', '树形图 + 力导向图，SVG 渲染'],
        ['数据格式', 'JSON (静态文件)', 'hanzi-dict.json / hanzi-index.json 等'],
        ['构建工具', 'Vite 5', '快速 HMR + 生产构建'],
        ['部署平台', 'GitHub Pages', '静态站点托管，免费 CDN'],
        ['核心数据量', '~500KB (gzip)', '3G 网络加载约 3-5 秒'],
        ['英汉索引', '~15MB (懒加载)', '首次英文搜索时按需加载'],
        ['性能优化', '懒加载 + 索引预计算\n+ 节点数量硬上限', '核心查询 <50ms\n可视化渲染 <200ms'],
    ],
    col_widths=[3.0, 5.0, 8.0],
    header_color=C_DARK
)

# === 3.5 后台数据库规模 ===
add_title('3.5 后台数据库规模', level=2)

# ---- Headline stats dashboard ----
add_para('核心数据指标', bold=True, size=Pt(11))

headline = new_table(1, 5, col_widths=[2.8, 2.8, 2.8, 2.8, 2.8])
h_data = [
    ('8,025\n收录汉字', C_RED),
    ('7,527\n字源标注', C_ORANGE),
    ('15,621\n字形图片', C_BLUE),
    ('30,208\n英文关键词', C_GOLD),
    ('207,658\n英汉词条', C_GREEN),
]
for i, (text, color) in enumerate(h_data):
    cell = headline.rows[0].cells[i]
    set_cell_bg(cell, color)
    set_cell_text(cell, text, bold=True, size=Pt(11), color=RGBColor(0xFF,0xFF,0xFF))

spacer()

# ---- Categorized data breakdown ----
add_para('数据资产分类明细', bold=True, size=Pt(11))

make_info_table(
    ['数据类别', '文件/来源', '规模', '说明'],
    [
        ['核心字库', 'hanzi-dict.json', '8,025 字', '全覆盖 CJK 基本汉字，含六书分类、IDS拆解、拼音、释义'],
        ['扩展字库', 'hanzi-data.json', '14,619 字', '含生僻字、异体字在内的扩展字符集'],
        ['《说文》数据', 'shuowen.json', '9,003 条', '每字含六书分类、部首、反切注音、释义'],
        ['字形图片库', 'glyphs/ 目录', '15,621 张 SVG', '甲骨文 1,533 + 金文 2,857 + 篆书 7,128 + 隶书 4,103'],
        ['英汉词典', 'en-word-index.json', '30,208 关键词\n207,658 词条', '基于 CC-CEDICT，支持英文反查中文'],
        ['通用规范字表', 'common-chars.json', '6,497 字', '覆盖一级 3,500 + 二级 ~3,000 常用字'],
        ['部件索引', 'hanzi-index.json', '1,597 个部件\n14,434 条引用', '含部件→汉字反向映射，最常用部件 氵 出现 427 次'],
        ['笔画数据', 'strokes.json', '1,018 字', '含笔顺动画数据'],
        ['文化数据', 'cultural.json', '1,002 字', '含字形演变说明、典故、词语示例'],
        ['简繁映射', 'simp-trad-map.json', '79 组配对', '简繁一对一/一对多映射关系'],
    ],
    col_widths=[2.5, 3.0, 3.5, 6.5],
    header_color=C_DARK
)

spacer()

# ---- Six-Book distribution ----
add_para('六书分类分布', bold=True, size=Pt(11))

make_info_table(
    ['六书类型', '数量 (hanzi-dict)', '数量 (说文)', '占比', '说明'],
    [
        ['形声字', '5,756', '5,186', '71.7%', '占绝对多数，声旁+形旁组合造字'],
        ['会意字', '1,564', '706', '19.5%', '多部件组合表意'],
        ['象形字', '207', '296', '2.6%', '图画式描摹实物'],
        ['指事字', '0', '57', '<1%', '抽象符号指示概念'],
        ['假借字', '—', '152', '—', '借用同音字表意'],
        ['其他/未分类', '498', '2,606', '6.2%', '含转注及暂未归类字'],
        ['合计', '8,025', '9,003', '100%', '两数据源六书体系互补'],
    ],
    col_widths=[2.2, 3.0, 2.5, 2.0, 5.8],
    header_color=C_ORANGE
)

spacer()

# ---- Manual annotation scale ----
add_para('人工标注数据', bold=True, size=Pt(11))

make_info_table(
    ['标注类型', '数量', '覆盖对象', '用途'],
    [
        ['变形偏旁标注', '23 组 + 2 特殊', '忄扌氵犭灬亻阝钅饣纟 等', '辅助拆解引擎还原部件原始形态'],
        ['幽灵部件标注', '15 个高频字', '买尽专长书为东乐头发后里农龙万', '标记简化字无造字意义的部件'],
        ['反义对标注', '59 组', '上下·大小·多少·高低·长短 等', '系联网络反义关系识别'],
        ['月/肉二义性', '关键词匹配规则', '身体词汇 → 肉; 天文词汇 → 月', '消除多义偏旁歧义'],
        ['阝左/右判定', '位置规则', '左阝→阜(山丘); 右阝→邑(城邑)', '消除同形偏旁歧义'],
    ],
    col_widths=[2.8, 3.0, 5.2, 4.5],
    header_color=C_GRAY
)

spacer()

# ---- In-memory index scale ----
add_para('内存索引运行时规模', bold=True, size=Pt(11))

make_info_table(
    ['索引结构', '键数量', '值类型', '内存占用（估算）', '平均查询'],
    [
        ['phoneticIndex', '~5,600 声旁', 'Set<汉字>', '~2 MB', '<1ms'],
        ['semanticIndex', '~5,700 形旁', 'Set<汉字>', '~2 MB', '<1ms'],
        ['pinyinIndex', '~1,200 音节', 'Set<汉字>', '~0.5 MB', '<1ms'],
        ['reverseIndex', '~1,600 部件', 'string[]', '~1 MB', '<1ms'],
        ['radicalIndex', '269 部首', 'string[]', '~0.5 MB', '<1ms'],
        ['structuralRank', '~1,600 部件', 'number', '~0.1 MB', '<1ms'],
        ['7 索引合计', '—', '—', '~6 MB', '全部 <1ms'],
    ],
    col_widths=[2.8, 3.0, 3.0, 3.5, 3.2],
    header_color=C_BLUE
)

# ================================================================
# SAVE
# ================================================================
doc.save(OUTPUT_PATH)
print(f'Done! Editable document saved to: {OUTPUT_PATH}')
print(f'File size: {os.path.getsize(OUTPUT_PATH) / 1024:.1f} KB')
