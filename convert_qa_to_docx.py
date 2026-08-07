"""Convert 答辩模拟问答.md to Word document."""
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_code_block(doc, lines):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(0.5)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '1A1A18')
        shd.set(qn('w:val'), 'clear')
        pPr.append(shd)
        r = p.add_run(line)
        r.font.name = 'Consolas'
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0xF5, 0xF0, 0xE8)

def add_inline_para(doc, text, indent=0, font_size=11, color=None):
    """Add paragraph with bold and code inline formatting."""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    parts = re.split(r'(\*\*.+?\*\*|`.+?`)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**') and len(part) > 4:
            r = p.add_run(part[2:-2])
            r.bold = True
        elif part.startswith('`') and part.endswith('`'):
            r = p.add_run(part[1:-1])
            r.font.name = 'Consolas'
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0xC2, 0x3B, 0x2A)
        else:
            r = p.add_run(part)
        r.font.size = Pt(font_size)
        r.font.name = '微软雅黑'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        if color:
            r.font.color.rgb = color

def add_table(doc, lines, start_idx):
    tbl_lines = []
    j = start_idx
    while j < len(lines) and lines[j].strip().startswith('|'):
        tbl_lines.append(lines[j].strip())
        j += 1
    if len(tbl_lines) < 2:
        return j
    header = [c.strip() for c in tbl_lines[0].split('|')[1:-1]]
    data = []
    for tl in tbl_lines[2:]:
        cells = [c.strip() for c in tl.split('|')[1:-1]]
        data.append(cells)
    if not header:
        return j
    t = doc.add_table(rows=1+len(data), cols=len(header))
    t.style = 'Light Grid Accent 1'
    for ci, hc in enumerate(header):
        cell = t.rows[0].cells[ci]
        p = cell.paragraphs[0]
        r = p.add_run(hc)
        r.bold = True
        r.font.size = Pt(10)
        r.font.name = '微软雅黑'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), 'C23B2A')
        shd.set(qn('w:val'), 'clear')
        cell._tc.get_or_add_tcPr().append(shd)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for ri, row in enumerate(data):
        for ci, tc in enumerate(row):
            if ci < len(header):
                cell = t.rows[ri+1].cells[ci]
                p = cell.paragraphs[0]
                r = p.add_run(tc)
                r.font.size = Pt(9.5)
                r.font.name = '微软雅黑'
                r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    doc.add_paragraph('')
    return j

def process_md_file(doc, md_path):
    """Process a single markdown file and add content to the document."""
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.read().split('\n')

    i = 0
    in_code = False
    code_buf = []

    while i < len(lines):
        line = lines[i]

        if line.strip().startswith('```'):
            if in_code:
                if code_buf:
                    add_code_block(doc, code_buf)
                code_buf = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        if line.strip() == '---':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), 'C23B2A')
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        if line.startswith('# ') and not line.startswith('## '):
            h = doc.add_heading(line[2:].strip(), level=1)
            for run in h.runs:
                run.font.name = '微软雅黑'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x18)
            i += 1
            continue

        if line.startswith('## '):
            h = doc.add_heading(line[3:].strip(), level=2)
            for run in h.runs:
                run.font.name = '微软雅黑'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                run.font.color.rgb = RGBColor(0xC2, 0x3B, 0x2A)
            i += 1
            continue

        if line.startswith('### '):
            h = doc.add_heading(line[4:].strip(), level=3)
            for run in h.runs:
                run.font.name = '微软雅黑'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                run.font.color.rgb = RGBColor(0x2D, 0x5F, 0x8A)
            i += 1
            continue

        if line.strip().startswith('|'):
            i = add_table(doc, lines, i)
            continue

        # Q header: **Q1：...**
        if re.match(r'^\*\*Q\d+[：:】\)]', line):
            text = line.replace('**', '')
            p = doc.add_paragraph()
            r = p.add_run(text)
            r.bold = True
            r.font.size = Pt(12)
            r.font.name = '微软雅黑'
            r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            r.font.color.rgb = RGBColor(0xC2, 0x3B, 0x2A)
            i += 1
            continue

        # Q header variant: ### Q1
        if re.match(r'^###\s*Q\d+', line):
            text = re.sub(r'^###\s*', '', line)
            p = doc.add_paragraph()
            r = p.add_run(text)
            r.bold = True
            r.font.size = Pt(12)
            r.font.name = '微软雅黑'
            r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            r.font.color.rgb = RGBColor(0xC2, 0x3B, 0x2A)
            i += 1
            continue

        # A header: **答**： or **A：**
        if re.match(r'^\*\*[答A]\*\*[：:]', line):
            text = line.replace('**', '')
            p = doc.add_paragraph()
            r = p.add_run(text)
            r.bold = True
            r.font.size = Pt(11)
            r.font.name = '微软雅黑'
            r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x18)
            i += 1
            continue

        # List items
        if re.match(r'^[-*] |^\d+\. ', line):
            text = re.sub(r'^[-*] |^\d+\. ', '', line)
            add_inline_para(doc, text, indent=0.8)
            i += 1
            continue

        # Bold headers like **第一，字源学深度**
        if line.strip().startswith('**') and line.strip().endswith('**') and len(line.strip()) < 100:
            text = line.strip().strip('*')
            p = doc.add_paragraph()
            r = p.add_run(text)
            r.bold = True
            r.font.size = Pt(11)
            r.font.name = '微软雅黑'
            r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            i += 1
            continue

        # Empty
        if not line.strip():
            i += 1
            continue

        # Regular paragraph
        add_inline_para(doc, line)
        i += 1


def main():
    import os

    # Find all markdown Q&A files
    md_files = []
    for f in os.listdir('.'):
        if f.endswith('.md') and '答辩' in f:
            md_files.append(f)
    md_files.sort()

    for md_file in md_files:
        doc = Document()
        section = doc.sections[0]
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)

        style = doc.styles['Normal']
        style.font.name = '微软雅黑'
        style.font.size = Pt(11)
        style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

        print(f'Processing: {md_file}')
        process_md_file(doc, md_file)

        docx_name = md_file.replace('.md', '.docx')
        doc.save(docx_name)
        print(f'Done! Saved to {docx_name}')

if __name__ == '__main__':
    main()
