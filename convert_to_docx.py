"""Convert 项目技术文档.md to a formatted Word document."""
import re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_table(doc, lines, start_idx):
    """Parse a markdown table and add it to the document."""
    # Collect all table lines
    table_lines = []
    i = start_idx
    while i < len(lines) and lines[i].strip().startswith('|'):
        table_lines.append(lines[i].strip())
        i += 1

    if len(table_lines) < 2:
        return i

    # Parse header
    header_cells = [c.strip() for c in table_lines[0].split('|')[1:-1]]
    # Skip separator line (index 1)
    # Parse data rows
    data_rows = []
    for line in table_lines[2:]:
        cells = [c.strip() for c in line.split('|')[1:-1]]
        data_rows.append(cells)

    # Create table
    table = doc.add_table(rows=1 + len(data_rows), cols=len(header_cells))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    for j, text in enumerate(header_cells):
        cell = table.rows[0].cells[j]
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, 'C23B2A')
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Data
    for i_row, row in enumerate(data_rows):
        for j, text in enumerate(row):
            if j < len(header_cells):
                cell = table.rows[i_row + 1].cells[j]
                p = cell.paragraphs[0]
                run = p.add_run(text)
                run.font.size = Pt(9.5)
                run.font.name = '微软雅黑'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph('')  # spacing after table
    return i

def parse_inline_formatting(paragraph, text):
    """Parse inline formatting within a text segment and add runs to paragraph."""
    # Handle bold (**...**) and inline code (`...`)
    # Split by patterns
    pattern = r'(\*\*\*.+?\*\*\*|\*\*.+?\*\*|`.+?`|\[.+?\]\(.+?\))'
    parts = re.split(pattern, text)

    for part in parts:
        if not part:
            continue

        # Bold+Italic ***...***
        if part.startswith('***') and part.endswith('***'):
            run = paragraph.add_run(part[3:-3])
            run.bold = True
            run.italic = True
            run.font.size = Pt(11)
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

        # Bold **...**
        elif part.startswith('**') and part.endswith('**') and len(part) > 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.size = Pt(11)
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

        # Inline code `...`
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xC2, 0x3B, 0x2A)

        # Link [...](...)
        elif part.startswith('[') and '](' in part:
            m = re.match(r'\[(.+?)\]\((.+?)\)', part)
            if m:
                run = paragraph.add_run(m.group(1))
                run.font.color.rgb = RGBColor(0x2D, 0x5F, 0x8A)
                run.underline = True
                run.font.size = Pt(11)
                run.font.name = '微软雅黑'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

        else:
            run = paragraph.add_run(part)
            run.font.size = Pt(11)
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def add_formatted_paragraph(doc, text, style=None, font_size=11, bold=False, italic=False, color=None, alignment=None):
    """Add a paragraph with mixed inline formatting."""
    p = doc.add_paragraph()
    if style:
        p.style = style

    parse_inline_formatting(p, text)

    if alignment is not None:
        p.alignment = alignment

    return p

def add_code_block(doc, code_lines):
    """Add a code block with dark background."""
    for line in code_lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(0.5)
        # Set paragraph shading
        pPr = p._p.get_or_add_pPr()
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), '1A1A18')
        shading.set(qn('w:val'), 'clear')
        pPr.append(shading)

        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xF5, 0xF0, 0xE8)
    doc.add_paragraph('')

def main():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # Read markdown
    with open('项目技术文档.md', 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_lines = []

    while i < len(lines):
        line = lines[i]

        # Code block start/end
        if line.strip().startswith('```'):
            if in_code_block:
                if code_lines:
                    add_code_block(doc, code_lines)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Horizontal rule
        if line.strip() == '---':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), 'C23B2A')
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # Headings
        if line.startswith('# ') and not line.startswith('## '):
            text = line[2:].strip()
            h = doc.add_heading(text, level=1)
            for run in h.runs:
                run.font.name = '微软雅黑'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x18)
            i += 1
            continue

        if line.startswith('## '):
            text = line[3:].strip()
            h = doc.add_heading(text, level=2)
            for run in h.runs:
                run.font.name = '微软雅黑'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                run.font.color.rgb = RGBColor(0xC2, 0x3B, 0x2A)
            i += 1
            continue

        if line.startswith('### '):
            text = line[4:].strip()
            h = doc.add_heading(text, level=3)
            for run in h.runs:
                run.font.name = '微软雅黑'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                run.font.color.rgb = RGBColor(0x2D, 0x5F, 0x8A)
            i += 1
            continue

        if line.startswith('#### '):
            text = line[5:].strip()
            h = doc.add_heading(text, level=4)
            for run in h.runs:
                run.font.name = '微软雅黑'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            i += 1
            continue

        # Table
        if line.strip().startswith('|'):
            i = add_table(doc, lines, i)
            continue

        # List items
        if re.match(r'^[-*] ', line):
            text = re.sub(r'^[-*] ', '', line)
            p = doc.add_paragraph(style='List Bullet')
            p.clear()
            parse_inline_formatting(p, text)
            i += 1
            continue

        if re.match(r'^\d+\. ', line):
            text = re.sub(r'^\d+\. ', '', line)
            p = doc.add_paragraph(style='List Number')
            p.clear()
            parse_inline_formatting(p, text)
            i += 1
            continue

        # Tree/sub-list items (indented with ├── └── │)
        if re.match(r'^\s*[├└│]', line):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.5)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            parse_inline_formatting(p, line)
            i += 1
            continue

        # Empty line
        if not line.strip():
            i += 1
            continue

        # Regular paragraph
        add_formatted_paragraph(doc, line)
        i += 1

    # Save
    output_path = '项目技术文档.docx'
    doc.save(output_path)
    print(f'Done! Saved to {output_path}')

if __name__ == '__main__':
    main()
